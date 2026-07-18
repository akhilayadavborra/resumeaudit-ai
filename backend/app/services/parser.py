"""
Extracts raw text from uploaded resume files (PDF or DOCX).
"""
import io
import fitz  # PyMuPDF
import docx
from fastapi import HTTPException


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception:
        raise HTTPException(status_code=400, detail="Unable to read PDF file. It may be corrupted.")

    text_parts = [page.get_text("text") for page in doc]
    doc.close()

    text = "\n".join(text_parts).strip()
    if not text:
        raise HTTPException(
            status_code=422,
            detail="No extractable text found. This PDF may be a scanned image without OCR text.",
        )
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception:
        raise HTTPException(status_code=400, detail="Unable to read DOCX file. It may be corrupted.")

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    text = "\n".join(parts).strip()
    if not text:
        raise HTTPException(status_code=422, detail="No extractable text found in this DOCX file.")
    return text


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Please upload a PDF or DOCX file.")