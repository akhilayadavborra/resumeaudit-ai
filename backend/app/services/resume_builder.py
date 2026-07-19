"""
Generates a clean, ATS-friendly resume file (DOCX or PDF) from structured
user input. Formatting follows ATS best practices: single-column layout,
standard section headers, no tables/text boxes/graphics, plain bullet points.
"""
import io
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from app.services.keyword_extractor import extract_skills
from app.services.formatting_checker import GENERIC_FILLER_EXAMPLES, STRONG_VERBS


def check_bullet_quality(bullet: str) -> dict:
    """
    Flags a single bullet as weak/generic (reusing the same rules as the
    analyzer) so the builder can warn the user before they even generate
    the file, instead of only finding out after the fact.
    """
    lower = bullet.lower()
    for entry in GENERIC_FILLER_EXAMPLES:
        if any(p in lower for p in entry["patterns"]):
            return {"is_weak": True, "reason": entry["problem"], "example_fix": entry["example_after"]}

    has_strong_verb = any(re.match(r"^" + v, lower) for v in STRONG_VERBS)
    has_metric = bool(re.search(r"\b\d+%|\$\d+|\b\d{2,}\b", bullet))

    if not has_strong_verb:
        return {"is_weak": True, "reason": "Doesn't start with a strong action verb.", "example_fix": None}
    if not has_metric:
        return {"is_weak": True, "reason": "No measurable result (number, %, or $).", "example_fix": None}

    return {"is_weak": False, "reason": None, "example_fix": None}


def _prioritize_skills(skills: list, job_description: str = None) -> list:
    """Reorders skills so ones matching the job description appear first."""
    if not job_description or len(job_description.strip()) < 20:
        return skills

    jd_skills = set(extract_skills(job_description))
    matching = [s for s in skills if s.lower() in jd_skills]
    non_matching = [s for s in skills if s.lower() not in jd_skills]
    return matching + non_matching


def generate_docx_resume(data) -> bytes:
    """Builds a single-column, ATS-parseable .docx resume."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Header: name + contact line
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(data.full_name)
    name_run.bold = True
    name_run.font.size = Pt(20)

    contact_parts = [p for p in [data.email, data.phone, data.linkedin] if p]
    contact_para = doc.add_paragraph(" | ".join(contact_parts))
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_section_header(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(3)

        # Adds a real horizontal divider line under the header (a bottom
        # border on the paragraph itself) — the standard professional
        # resume convention, instead of just relying on blank spacing.
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom_border = OxmlElement("w:bottom")
        bottom_border.set(qn("w:val"), "single")
        bottom_border.set(qn("w:sz"), "8")
        bottom_border.set(qn("w:space"), "1")
        bottom_border.set(qn("w:color"), "444444")
        p_bdr.append(bottom_border)
        p_pr.append(p_bdr)

    if data.summary:
        add_section_header("Professional Summary")
        doc.add_paragraph(data.summary)

    if data.experience:
        add_section_header("Experience")
        for exp in data.experience:
            line = doc.add_paragraph()
            line.add_run(f"{exp.title} — {exp.company}").bold = True
            date_line = doc.add_paragraph(f"{exp.start_date} - {exp.end_date}")
            date_line.paragraph_format.space_after = Pt(2)
            for bullet in exp.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    if data.projects:
        add_section_header("Projects")
        for proj in data.projects:
            line = doc.add_paragraph()
            line.add_run(proj.name).bold = True
            if proj.description:
                doc.add_paragraph(proj.description)
            for bullet in proj.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    if data.education:
        add_section_header("Education")
        for edu in data.education:
            line = doc.add_paragraph()
            line.add_run(f"{edu.degree} — {edu.school}").bold = True
            if edu.start_date or edu.end_date:
                doc.add_paragraph(f"{edu.start_date} - {edu.end_date}")

    if data.skills:
        add_section_header("Skills")
        ordered_skills = _prioritize_skills(data.skills, data.job_description)
        doc.add_paragraph(", ".join(ordered_skills))

    if data.certifications:
        add_section_header("Certifications")
        for cert in data.certifications:
            doc.add_paragraph(cert, style="List Bullet")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_pdf_resume(data) -> bytes:
    """Builds a single-column, ATS-parseable .pdf resume."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch, leftMargin=0.75 * inch, rightMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()

    name_style = ParagraphStyle("NameStyle", parent=styles["Title"], fontSize=20, alignment=1, spaceAfter=4)
    contact_style = ParagraphStyle("ContactStyle", parent=styles["Normal"], alignment=1, fontSize=10, textColor="#444444")
    section_style = ParagraphStyle("SectionStyle", parent=styles["Heading2"], fontSize=12, spaceBefore=14, spaceAfter=2, textColor="#1a1a1a")
    body_style = ParagraphStyle("BodyStyle", parent=styles["BodyText"], fontSize=10, leading=14)
    bold_line_style = ParagraphStyle("BoldLine", parent=body_style, fontName="Helvetica-Bold")

    elements = [
        Paragraph(data.full_name, name_style),
        Paragraph(" | ".join([p for p in [data.email, data.phone, data.linkedin] if p]), contact_style),
    ]

    if data.summary:
        elements.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.75, color="#444444", spaceAfter=6))
        elements.append(Paragraph(data.summary, body_style))

    if data.experience:
        elements.append(Paragraph("EXPERIENCE", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.75, color="#444444", spaceAfter=6))
        for exp in data.experience:
            elements.append(Paragraph(f"{exp.title} — {exp.company}", bold_line_style))
            elements.append(Paragraph(f"{exp.start_date} - {exp.end_date}", body_style))
            if exp.bullets:
                elements.append(ListFlowable([ListItem(Paragraph(b, body_style)) for b in exp.bullets], bulletType="bullet"))
            elements.append(Spacer(1, 6))

    if data.projects:
        elements.append(Paragraph("PROJECTS", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.75, color="#444444", spaceAfter=6))
        for proj in data.projects:
            elements.append(Paragraph(proj.name, bold_line_style))
            if proj.description:
                elements.append(Paragraph(proj.description, body_style))
            if proj.bullets:
                elements.append(ListFlowable([ListItem(Paragraph(b, body_style)) for b in proj.bullets], bulletType="bullet"))
            elements.append(Spacer(1, 6))

    if data.education:
        elements.append(Paragraph("EDUCATION", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.75, color="#444444", spaceAfter=6))
        for edu in data.education:
            elements.append(Paragraph(f"{edu.degree} — {edu.school}", bold_line_style))
            if edu.start_date or edu.end_date:
                elements.append(Paragraph(f"{edu.start_date} - {edu.end_date}", body_style))

    if data.skills:
        elements.append(Paragraph("SKILLS", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.75, color="#444444", spaceAfter=6))
        ordered_skills = _prioritize_skills(data.skills, data.job_description)
        elements.append(Paragraph(", ".join(ordered_skills), body_style))

    if data.certifications:
        elements.append(Paragraph("CERTIFICATIONS", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.75, color="#444444", spaceAfter=6))
        elements.append(ListFlowable([ListItem(Paragraph(c, body_style)) for c in data.certifications], bulletType="bullet"))

    doc.build(elements)
    buffer.seek(0)
    return buffer.read()